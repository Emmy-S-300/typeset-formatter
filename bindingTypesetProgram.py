"""
This program is to take any of my favorite online stories and format
 them into a typset so that I can print them and bind them.
"""
import sys, os
from docx import Document
from bs4 import BeautifulSoup
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
#from tkinter import Tk, filedialog

#I want to import an html file
#this only reads and returns the raw html content.
def html_file_read(file_path):
    with open(file_path, "r", encoding="utf-8") as file:
        return file.read()

def set_landscape_and_margins(document):
    section = document.sections[0]    #define the section. There is usually only on according to https://python-docx.readthedocs.io/en/latest/user/sections.html

    section.orientation = WD_ORIENT.LANDSCAPE   #set the section to landscape mode

    #match the hieght and width to the landscape layout
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    section.orientation, section.page_width, section.page_height #(LANDSCAPE (1), 10058400, 7772400)

    #Adjust margins 
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

def get_metadata(soup):
    #A dictionary to record all of the metadata needed for a titlepage and summary page
    metadata = {
    "title": "...",
    "author": "...",
    "published": "...",
    "completed": None,
    "words": "...",
    "chapters": "...",
    "series": None,
    "summary": "..."
    }
    #Find the data and update the dictionary
    
    #title and author
    metadata["title"] = soup.find("div", {"class": "meta"}).find("h1").get_text(strip=True)
    metadata["author"] = soup.find("div", {"class": "byline"}).get_text(strip=True)
    
    #get the metadata from the stats div
    stats_dd = soup.find("dt", string="Stats:").find_next("dd").get_text(" ", strip=True)

  #if the published is in the stats block 
    if "Published:" in stats_dd:
        metadata["published"] = stats_dd.split("Published:")[1].split("Completed:")[0].strip()

    #if the completed exists
    if "Completed:" in stats_dd:
        metadata["completed"] = stats_dd.split("Completed:")[1].split("Words:")[0].strip()

    #Words
    if "Words:" in stats_dd:
        metadata["words"] = stats_dd.split("Words:")[1].split("Chapters:")[0].strip()

    #Chapters
    if "Chapters:" in stats_dd:
        metadata["chapters"] = stats_dd.split("Chapters:")[1].split()[0]
    
    #if there is a series the work is a part of
    series_dd = soup.find("dd", {"class": "series"})
    metadata["series"] = series_dd.get_text(" ", strip=True) if series_dd else None
    
    #if there is a Summary
    summary_block = soup.find("blockquote", {"class": "userstuff"})
    metadata["summary"] = summary_block.get_text(strip=True) if summary_block else None

    return metadata

def set_title_page(document, metadata):
    #add blank pages in the beginning
    for i in range(3):
        document.add_page_break()

    #center the title page vertically
    center_vert = document.add_paragraph()
    center_vert.paragraph_format.space_before = Pt(125)


    title_text =document.add_paragraph(metadata["title"]) #get the title
    format_font_size_alignment(title_text, font_name = 'Garamond', size = 28, bold = True, alignment = WD_ALIGN_PARAGRAPH.CENTER)  #apply the font and size to the title

    author_text = document.add_paragraph({ metadata["author"]})
    format_font_size_alignment(author_text, font_name = 'Garamond', size = 24, bold = True, alignment = WD_ALIGN_PARAGRAPH.CENTER)  #apply the font and size to the author
    
    #published
    if metadata["published"]:
        pub_text = document.add_paragraph(f"Published: {metadata['published']}")
        format_font_size_alignment(pub_text, font_name = 'Garamond', size = 14, bold = True, alignment = WD_ALIGN_PARAGRAPH.CENTER)  #apply the font and size to the published text
    
    #completed
    if metadata["completed"]:
        comp_text = document.add_paragraph(f"Completed: {metadata['completed']}")
        format_font_size_alignment(comp_text, font_name = 'Garamond', size = 14, bold = True, alignment = WD_ALIGN_PARAGRAPH.CENTER)  #apply the font and size to the completed text
    
    #words
    if metadata["words"]:
        words_text = document.add_paragraph(f"Words: {metadata["words"]}")
        format_font_size_alignment(words_text, font_name = 'Garamond', size = 14, bold = True, alignment = WD_ALIGN_PARAGRAPH.CENTER)  #apply the font and size to the words_text
    
    #chapters
    if metadata["chapters"]:
        chapters_text = document.add_paragraph(f"Chapters: {metadata["chapters"]}")
        format_font_size_alignment(chapters_text, font_name = 'Garamond', size = 14, bold = True, alignment = WD_ALIGN_PARAGRAPH.CENTER)  #apply the font and size to the chapters_text
    
    #series
    if metadata["series"]:
        series_text = document.add_paragraph(f"Series: {metadata["series"]}")
        format_font_size_alignment(series_text, font_name = 'Garamond', size = 14, bold = True, alignment = WD_ALIGN_PARAGRAPH.CENTER)  #apply the font and size to the series_text.

    #binding info
    binding_info = document.add_paragraph(
        f"\n "
        f"Typesetting & Bookbinding by:\n"
        f"Add your name here"
        )
    format_font_size_alignment(binding_info, font_name = 'Garamond', size = 14, bold = True, alignment = WD_ALIGN_PARAGRAPH.CENTER)  #apply the font and size to the series_text.


    document.add_page_break()

def add_summary_page(document, metadata):
    #Summary
    if metadata["summary"]:
        # add "summary" for the heading on the page.
        summary_heading = document.add_paragraph("Summary:")
        format_font_size_alignment(summary_heading, font_name = 'Garamond', size = 24, bold = True, alignment = WD_ALIGN_PARAGRAPH.CENTER)  #apply the font and size to the series_text.
    
        #add the summary text
        summary_text = document.add_paragraph(metadata["summary"])
        format_font_size_alignment(summary_text, font_name = 'Garamond', size = 14, bold = False, alignment = WD_ALIGN_PARAGRAPH.JUSTIFY)  #apply the font and size to the series_text.

    document.add_page_break()
    
def format_font_size_alignment(paragraph, font_name = 'Garamond', size = 11, bold = False, alignment = WD_ALIGN_PARAGRAPH.JUSTIFY, indent = False):
   #ensure a run happens
   if paragraph.runs:
       run = paragraph.runs[0]
   else:
       run = paragraph.add_run()

       
   font = run.font
   font.name = font_name
   font.size = Pt(size)
   font.bold = bold
   paragraph.alignment = alignment

   # if there is an indent needed
   if indent == True: # if there is an indent needed
      
       paragraph_format = paragraph.paragraph_format    #set the format to the variable
       paragraph_format.first_line_indent = Inches(0.5)

def add_page_numbers(section):
    footer = section.footer # create the footer section
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER # center the footers

    run = paragraph.add_run()




#I want the html file to be written to a word document
def write_to_word_doc(diff_content_blocks, output_path, metadata):
    #creating variables
    chap_1 = True
    
    document = Document()   #create a new document

    set_landscape_and_margins(document)# calling the landsacpe function
    set_title_page(document, metadata)# calling the title page function
    add_summary_page(document, metadata)    #calling the summary page function

    new_section = document.add_section(WD_SECTION.NEW_PAGE) #creating a new section so that page numbers start on the story and not the title/summary
    add_page_numbers(new_section)


    for blocks in diff_content_blocks:
        text = blocks.get_text()

        if blocks.name == "h1":     # it it is a title h1 tag
            if chap_1:
                #make sure that it starts on the right page. That being gthe odd pages. This will [ut it on the right side of the book when it is opened.
                document.add_page_break()
                chap_1 = False
            else:
                document.add_page_break()

            heading = document.add_paragraph(text) #get the text
            format_font_size_alignment(heading, font_name = 'Garamond', size = 24, bold = True, alignment = WD_ALIGN_PARAGRAPH.CENTER)  #apply the font and size to the heading

        elif blocks.name == 'h2':   # if it is a title h2 tag
            document.add_page_break()   #start on a new page.
            heading = document.add_paragraph(text) #get the text
            format_font_size_alignment(heading, font_name = 'Garamond', size = 18, bold = True, alignment = WD_ALIGN_PARAGRAPH.CENTER)  #apply the font and size to the heading 2

        elif blocks.name == 'p':    # if it is a paragraph tag
            paragraph = document.add_paragraph(text) #get the text
            format_font_size_alignment(paragraph, font_name = 'Garamond', size = 11, bold = True, alignment = WD_ALIGN_PARAGRAPH.JUSTIFY, indent = True)  #apply the font and size to the paragraph text

    document.save(output_path)  #save the text in the output in the document
    return document

#where all the magic happens
if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python your_script.py <html_file>")
    else:
        html_path = sys.argv[1]
        html_content = html_file_read(html_path) #get the raw html
        soup = BeautifulSoup(html_content, 'html.parser') #use beautiful soup to parse the html

       #This is the filter
       #find the chapters in the html
        story_divs = soup.find("div", {"id": "chapters"})
        #find the paragraphs in the chapters
        if story_divs:
            diff_content_blocks = story_divs.find_all(["h1", "h2","p"]) #find the headings 11 headings 2, and paragraphs
            # for p in paragraphs: # for the paragraphs found
            #        print(p.get_text())#print all the found paragraphs
        else:
            print("could not find story content! :(")
            diff_content_blocks = [] # a fallback
         
        #write to the word doc
        base_name = input("Title the typeset: ")
        output_path = base_name + ".docx"    #force a document file type
        counter = 1

        #ensurer that files are not over written.
        while os.path.exists(output_path):
            output_path = f"{base_name}_{counter}.docx"
            counter +=1
       
        metadata= get_metadata(soup)

        write_to_word_doc(diff_content_blocks, output_path, metadata)
        print(f"HTML content has been written to {output_path}")    #the f string will make the actual name of the file show up

"""
Now I want the pages to fold like a book too.
I also want a separate title page.
picture formatting 
book leftlet settings
comments section at the very end
blank oages added at the beginning and end
Ability to take user input and and name the file output.
page breaks before chapters
maybe use selenium to go in and fix the custom stuff python -docx can't ie bookfold.
"""
